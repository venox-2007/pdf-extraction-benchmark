"""Generate DOCX and PDF exports of docs/final_benchmark_report.md.

Parses the markdown report into a small structured block representation
(title, headings, paragraphs, bullet lists, tables, code blocks, page-break
markers) and renders that single parsed structure into both output formats,
so the two exports never drift from each other or from the markdown source.

This script only reads docs/final_benchmark_report.md and writes
docs/final_benchmark_report.docx / .pdf — it does not touch any benchmark
code, benchmark output, or other documentation.
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "docs" / "final_benchmark_report.md"
DOCX_PATH = ROOT / "docs" / "final_benchmark_report.docx"
PDF_PATH = ROOT / "docs" / "final_benchmark_report.pdf"


# ---------------------------------------------------------------------------
# Markdown parsing
# ---------------------------------------------------------------------------


@dataclass(slots=True)
class Block:
    kind: str  # title|meta|toc_skip|h2|h3|para|bullets|table|code|hr
    text: str = ""
    items: list[str] = field(default_factory=list)
    rows: list[list[str]] = field(default_factory=list)


def parse_markdown(md_text: str) -> list[Block]:
    lines = md_text.splitlines()
    blocks: list[Block] = []
    i = 0
    n = len(lines)
    in_toc = False
    in_code = False
    code_buf: list[str] = []

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if in_code:
            if stripped.startswith("```"):
                blocks.append(Block(kind="code", text="\n".join(code_buf)))
                code_buf = []
                in_code = False
            else:
                code_buf.append(line)
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        if stripped.startswith("```"):
            in_code = True
            i += 1
            continue

        if stripped == "---":
            blocks.append(Block(kind="hr"))
            in_toc = False
            i += 1
            continue

        if stripped.startswith("# "):
            blocks.append(Block(kind="title", text=stripped[2:].strip()))
            i += 1
            continue

        if stripped.startswith("### Evaluating"):
            blocks.append(Block(kind="subtitle", text=stripped[4:].strip()))
            i += 1
            continue

        if stripped.startswith("## Table of Contents"):
            in_toc = True
            i += 1
            continue

        if in_toc:
            # Skip the manually-written numbered TOC list; both exports
            # generate a real, page-numbered TOC instead.
            i += 1
            continue

        if stripped.startswith("## "):
            blocks.append(Block(kind="h2", text=stripped[3:].strip()))
            i += 1
            continue

        if stripped.startswith("### "):
            blocks.append(Block(kind="h3", text=stripped[4:].strip()))
            i += 1
            continue

        meta_pattern = r"^\*\*(Project|Author|Supervisor|Duration|Document type):"
        if stripped.startswith("**") and re.match(meta_pattern, stripped):
            blocks.append(Block(kind="meta", text=stripped))
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines = []
            while i < n and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = [
                [cell.strip() for cell in row.strip("|").split("|")]
                for row in table_lines
                if not re.match(r"^\|?[\s:|-]+\|?$", row)
            ]
            blocks.append(Block(kind="table", rows=rows))
            continue

        if stripped.startswith("- ") or stripped.startswith("-  "):
            items = []
            while i < n:
                is_new_item = lines[i].strip().startswith("- ")
                is_wrapped_line = lines[i].startswith("  ") and items
                if not (is_new_item or is_wrapped_line):
                    break
                if is_new_item:
                    items.append(lines[i].strip()[2:].strip())
                else:
                    items[-1] += " " + lines[i].strip()
                i += 1
            blocks.append(Block(kind="bullets", items=items))
            continue

        # Paragraph: accumulate until blank line / next structural marker
        para_lines = [stripped]
        i += 1
        while i < n and lines[i].strip() and not lines[i].strip().startswith(
            ("#", "|", "- ", "```", "---")
        ):
            para_lines.append(lines[i].strip())
            i += 1
        blocks.append(Block(kind="para", text=" ".join(para_lines)))

    return blocks


def inline_to_html(text: str) -> str:
    """Convert the markdown inline subset used in this report to mini-HTML
    understood by reportlab Paragraph (and re-parsed for docx runs)."""
    text = text.replace("&", "&amp;")
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"(?<!\*)\*([^*]+?)\*(?!\*)", r"<i>\1</i>", text)
    text = re.sub(r"`([^`]+?)`", r'<font face="DejaVuSansMono">\1</font>', text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)  # drop markdown links, keep label
    return text


_META_RE = re.compile(r"^\*\*(.+?):\*\*\s*(.*)$")


def split_meta(text: str) -> tuple[str, str]:
    """Split a '**Label:** value' meta line into (label, value)."""
    match = _META_RE.match(text.strip())
    if match:
        return match.group(1).strip(), match.group(2).strip()
    label, _, value = text.strip("*").partition(":")
    return label.strip(), value.strip()


_TAG_RE = re.compile(r"<(/?)(b|i|font)(?:\s+face=\"([^\"]*)\")?>")


def html_runs(html_text: str) -> list[tuple[str, dict]]:
    """Split mini-HTML into (text, {bold, italic, mono}) runs for docx."""
    runs: list[tuple[str, dict]] = []
    state = {"bold": False, "italic": False, "mono": False}
    pos = 0
    for m in _TAG_RE.finditer(html_text):
        if m.start() > pos:
            runs.append((html_text[pos : m.start()], dict(state)))
        closing, tag, _face = m.groups()
        if tag == "b":
            state["bold"] = not closing
        elif tag == "i":
            state["italic"] = not closing
        elif tag == "font":
            state["mono"] = not closing
        pos = m.end()
    if pos < len(html_text):
        runs.append((html_text[pos:], dict(state)))
    return [(t, s) for t, s in runs if t]


# ---------------------------------------------------------------------------
# DOCX rendering
# ---------------------------------------------------------------------------


def render_docx(blocks: list[Block]) -> None:
    from docx import Document
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches, Pt

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for level, size in ((1, 20), (2, 16), (3, 13)):
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True

    def add_page_number_field(paragraph) -> None:
        run = paragraph.add_run()
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        run._r.append(fld)

    # --- Footer with page numbers ---
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_para.add_run("Page ")
    add_page_number_field(footer_para)
    footer_para.add_run(" of ")
    run = footer_para.add_run()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "NUMPAGES")
    run._r.append(fld)

    # --- Title page ---
    title_block = next(b for b in blocks if b.kind == "title")
    subtitle_block = next((b for b in blocks if b.kind == "subtitle"), None)
    meta_blocks = [b for b in blocks if b.kind == "meta"]

    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph(title_block.text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    p.runs[0].font.size = Pt(22)
    p.runs[0].font.bold = True
    p.runs[0].font.color.rgb = None

    if subtitle_block:
        p = doc.add_paragraph(subtitle_block.text)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(18)
        p.runs[0].font.size = Pt(15)
        p.runs[0].font.italic = True

    doc.add_paragraph()
    for mb in meta_blocks:
        label, value = split_meta(mb.text)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        p.add_run(value)

    doc.add_page_break()

    # --- Table of contents (native Word TOC field; updates on F9 / open) ---
    doc.add_heading("Table of Contents", level=1)
    toc_para = doc.add_paragraph()
    run = toc_para.add_run()
    fc_begin = OxmlElement("w:fldChar")
    fc_begin.set(qn("w:fldCharType"), "begin")
    r_begin = OxmlElement("w:r")
    r_begin.append(fc_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    r_instr = OxmlElement("w:r")
    r_instr.append(instr)
    fc_sep = OxmlElement("w:fldChar")
    fc_sep.set(qn("w:fldCharType"), "separate")
    r_sep = OxmlElement("w:r")
    r_sep.append(fc_sep)
    t = OxmlElement("w:t")
    t.text = "Right-click and select Update Field to populate this table of contents."
    r_text = OxmlElement("w:r")
    r_text.append(t)
    fc_end = OxmlElement("w:fldChar")
    fc_end.set(qn("w:fldCharType"), "end")
    r_end = OxmlElement("w:r")
    r_end.append(fc_end)
    for el in (r_begin, r_instr, r_sep, r_text, r_end):
        toc_para._p.append(el)
    doc.add_page_break()

    # --- Body ---
    table_counter = 0
    pending_break = False
    for block in blocks:
        if block.kind in ("title", "subtitle", "meta"):
            continue
        if block.kind == "hr":
            pending_break = True
            continue
        if block.kind == "h2":
            if pending_break:
                doc.add_page_break()
                pending_break = False
            doc.add_heading(block.text, level=1)
        elif block.kind == "h3":
            doc.add_heading(block.text, level=2)
        elif block.kind == "para":
            p = doc.add_paragraph()
            for text, fmt in html_runs(inline_to_html(block.text)):
                run = p.add_run(text)
                run.bold = fmt["bold"]
                run.italic = fmt["italic"]
                if fmt["mono"]:
                    run.font.name = "Consolas"
        elif block.kind == "bullets":
            for item in block.items:
                p = doc.add_paragraph(style="List Bullet")
                for text, fmt in html_runs(inline_to_html(item)):
                    run = p.add_run(text)
                    run.bold = fmt["bold"]
                    run.italic = fmt["italic"]
                    if fmt["mono"]:
                        run.font.name = "Consolas"
        elif block.kind == "code":
            p = doc.add_paragraph()
            run = p.add_run(block.text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        elif block.kind == "table":
            table_counter += 1
            rows = block.rows
            if not rows:
                continue
            table = doc.add_table(rows=len(rows), cols=len(rows[0]))
            table.style = "Light Grid Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for r, row in enumerate(rows):
                for c, cell_text in enumerate(row):
                    cell = table.rows[r].cells[c]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    for text, fmt in html_runs(inline_to_html(cell_text)):
                        run = p.add_run(text)
                        run.bold = fmt["bold"] or r == 0
                        run.italic = fmt["italic"]
                        run.font.size = Pt(9.5)
            # Keep header row from splitting across a page break.
            for row in table.rows:
                tr = row._tr
                trPr = tr.get_or_add_trPr()
                cant_split = OxmlElement("w:cantSplit")
                trPr.append(cant_split)
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = caption.add_run(f"Table {table_counter}")
            run.italic = True
            run.font.size = Pt(9)

    doc.save(str(DOCX_PATH))
    print(f"Wrote {DOCX_PATH}")


# ---------------------------------------------------------------------------
# PDF rendering
# ---------------------------------------------------------------------------


def render_pdf(blocks: list[Block]) -> None:
    import matplotlib
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.platypus import (
        BaseDocTemplate,
        NextPageTemplate,
        PageBreak,
        PageTemplate,
        Paragraph,
        Spacer,
        Table,
        TableStyle,
    )
    from reportlab.platypus.frames import Frame
    from reportlab.platypus.tableofcontents import TableOfContents

    font_dir = Path(matplotlib.__file__).resolve().parent / "mpl-data" / "fonts" / "ttf"
    pdfmetrics.registerFont(TTFont("DejaVuSans", str(font_dir / "DejaVuSans.ttf")))
    pdfmetrics.registerFont(TTFont("DejaVuSans-Bold", str(font_dir / "DejaVuSans-Bold.ttf")))
    pdfmetrics.registerFont(TTFont("DejaVuSans-Oblique", str(font_dir / "DejaVuSans-Oblique.ttf")))
    pdfmetrics.registerFont(
        TTFont("DejaVuSans-BoldOblique", str(font_dir / "DejaVuSans-BoldOblique.ttf"))
    )
    pdfmetrics.registerFont(TTFont("DejaVuSansMono", str(font_dir / "DejaVuSansMono.ttf")))
    pdfmetrics.registerFontFamily(
        "DejaVuSans",
        normal="DejaVuSans",
        bold="DejaVuSans-Bold",
        italic="DejaVuSans-Oblique",
        boldItalic="DejaVuSans-BoldOblique",
    )

    styles = getSampleStyleSheet()
    base = ParagraphStyle(
        "Body", parent=styles["Normal"], fontName="DejaVuSans", fontSize=10.5, leading=15,
        spaceAfter=8,
    )
    h1 = ParagraphStyle(
        "H1", parent=styles["Heading1"], fontName="DejaVuSans-Bold", fontSize=18,
        textColor=colors.HexColor("#1F3864"), spaceBefore=18, spaceAfter=10,
    )
    h2 = ParagraphStyle(
        "H2", parent=styles["Heading2"], fontName="DejaVuSans-Bold", fontSize=13.5,
        textColor=colors.HexColor("#2E5395"), spaceBefore=12, spaceAfter=6,
    )
    bullet_style = ParagraphStyle(
        "Bullet", parent=base, leftIndent=16, bulletIndent=4, spaceAfter=4,
    )
    code_style = ParagraphStyle(
        "Code", parent=base, fontName="DejaVuSansMono", fontSize=8.3, leading=11,
        backColor=colors.HexColor("#F2F2F2"), borderPadding=6, spaceAfter=10,
    )
    caption_style = ParagraphStyle(
        "Caption", parent=base, fontName="DejaVuSans-Oblique", fontSize=9,
        alignment=1, textColor=colors.grey, spaceAfter=14,
    )
    title_style = ParagraphStyle(
        "TitleBig", parent=base, fontName="DejaVuSans-Bold", fontSize=22, leading=27,
        alignment=1, textColor=colors.HexColor("#1F3864"), spaceAfter=16,
    )
    subtitle_style = ParagraphStyle(
        "SubtitleBig", parent=base, fontName="DejaVuSans-Oblique", fontSize=15, alignment=1,
        spaceAfter=24,
    )
    meta_style = ParagraphStyle("Meta", parent=base, alignment=1, spaceAfter=4)
    toc_h_style = ParagraphStyle(
        "TOCHeading1", fontName="DejaVuSans-Bold", fontSize=12, leftIndent=0, spaceAfter=6,
    )
    toc_sub_style = ParagraphStyle(
        "TOCHeading2", fontName="DejaVuSans", fontSize=10.5, leftIndent=18, spaceAfter=4,
    )

    def html_to_paragraph(text: str, style) -> Paragraph:
        html = inline_to_html(text)
        html = html.replace('<font face="DejaVuSansMono">', '<font face="DejaVuSansMono" size="9">')
        return Paragraph(html, style)

    def on_page(canvas, doc_):
        canvas.saveState()
        canvas.setFont("DejaVuSans", 8.5)
        canvas.setFillColor(colors.grey)
        page_num_text = f"Page {doc_.page}"
        canvas.drawCentredString(LETTER[0] / 2, 0.55 * inch, page_num_text)
        canvas.drawString(0.85 * inch, 0.55 * inch, "PDF Extraction Benchmark Report")
        canvas.restoreState()

    def on_title_page(canvas, doc_):
        canvas.saveState()
        canvas.setFont("DejaVuSans", 8.5)
        canvas.setFillColor(colors.grey)
        canvas.drawCentredString(LETTER[0] / 2, 0.55 * inch, "")
        canvas.restoreState()

    class ReportDocTemplate(BaseDocTemplate):
        """Notifies the TOC of each heading's page number as it is drawn."""

        def afterFlowable(self, flowable):
            if not isinstance(flowable, Paragraph):
                return
            text = flowable.getPlainText()
            style_name = flowable.style.name
            if style_name == "H1":
                level = 0
            elif style_name == "H2":
                level = 1
            else:
                return
            bookmark = getattr(flowable, "_bookmarkName", None)
            if bookmark:
                self.canv.bookmarkPage(bookmark)
                self.canv.addOutlineEntry(text, bookmark, level=level)
            self.notify("TOCEntry", (level, text, self.page, bookmark))

    doc_template = ReportDocTemplate(
        str(PDF_PATH),
        pagesize=LETTER,
        leftMargin=0.85 * inch,
        rightMargin=0.85 * inch,
        topMargin=0.9 * inch,
        bottomMargin=0.9 * inch,
        title="PDF Extraction Benchmark Report",
        author="Yug Agrawal",
    )
    frame = Frame(
        doc_template.leftMargin,
        doc_template.bottomMargin,
        doc_template.width,
        doc_template.height,
        id="normal",
    )
    doc_template.addPageTemplates(
        [
            PageTemplate(id="Title", frames=[frame], onPage=on_title_page),
            PageTemplate(id="Body", frames=[frame], onPage=on_page),
        ]
    )

    toc = TableOfContents()
    toc.levelStyles = [toc_h_style, toc_sub_style]

    story: list = []

    title_block = next(b for b in blocks if b.kind == "title")
    subtitle_block = next((b for b in blocks if b.kind == "subtitle"), None)
    meta_blocks = [b for b in blocks if b.kind == "meta"]

    story.append(Spacer(1, 2.2 * inch))
    story.append(Paragraph(title_block.text, title_style))
    if subtitle_block:
        story.append(Paragraph(subtitle_block.text, subtitle_style))
    for mb in meta_blocks:
        label, value = split_meta(mb.text)
        story.append(Paragraph(f"<b>{label}:</b> {value}", meta_style))

    story.append(NextPageTemplate("Body"))
    story.append(PageBreak())
    story.append(Paragraph("Table of Contents", h1))
    story.append(toc)
    story.append(PageBreak())

    heading_counter = [0, 0]
    table_counter = 0
    pending_break = False

    def make_heading(text: str, style, level: int) -> Paragraph:
        heading_counter[level - 1] += 1
        if level == 1:
            heading_counter[1] = 0
        key = f"h{level}-{heading_counter[0]}-{heading_counter[1]}"
        para = Paragraph(text, style)
        para._bookmarkName = key
        return para

    for block in blocks:
        if block.kind in ("title", "subtitle", "meta"):
            continue
        if block.kind == "hr":
            pending_break = True
            continue
        if block.kind == "h2":
            if pending_break or story:
                story.append(PageBreak())
                pending_break = False
            story.append(make_heading(block.text, h1, level=1))
        elif block.kind == "h3":
            story.append(make_heading(block.text, h2, level=2))
        elif block.kind == "para":
            story.append(html_to_paragraph(block.text, base))
        elif block.kind == "bullets":
            for item in block.items:
                story.append(html_to_paragraph(f"• {item}", bullet_style))
        elif block.kind == "code":
            escaped = block.text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            story.append(Paragraph(escaped.replace("\n", "<br/>"), code_style))
        elif block.kind == "table":
            table_counter += 1
            rows = block.rows
            if not rows:
                continue
            ncols = len(rows[0])
            avail_width = doc_template.width
            col_width = avail_width / ncols
            cell_style = ParagraphStyle(
                "Cell", parent=base, fontSize=8.3, leading=10.5, spaceAfter=0
            )
            header_style = ParagraphStyle(
                "CellHeader", parent=cell_style, fontName="DejaVuSans-Bold",
                textColor=colors.white,
            )
            data = []
            for r, row in enumerate(rows):
                style_for_row = header_style if r == 0 else cell_style
                data.append([html_to_paragraph(cell, style_for_row) for cell in row])
            tbl = Table(data, colWidths=[col_width] * ncols, repeatRows=1)
            row_bands = [colors.white, colors.HexColor("#F2F5FA")]
            tbl.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F3864")),
                        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#BFBFBF")),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), row_bands),
                        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                        ("TOPPADDING", (0, 0), (-1, -1), 4),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                        ("LEFTPADDING", (0, 0), (-1, -1), 5),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            story.append(tbl)
            story.append(Paragraph(f"Table {table_counter}", caption_style))

    doc_template.multiBuild(story)
    print(f"Wrote {PDF_PATH}")


def main() -> None:
    md_text = MD_PATH.read_text(encoding="utf-8")
    blocks = parse_markdown(md_text)
    render_docx(blocks)
    render_pdf(blocks)


if __name__ == "__main__":
    main()
